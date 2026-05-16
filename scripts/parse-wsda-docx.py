"""Extract impromptu motions from WSDA*.docx in repo root -> web/content/wsda-impromptu-parsed.json"""
from __future__ import annotations

import json
import re
from pathlib import Path

try:
    import docx  # type: ignore
except ImportError as e:
    raise SystemExit("pip install python-docx") from e

REPO = Path(__file__).resolve().parents[2]
WEB = REPO / "web"
OUT = WEB / "content" / "wsda-impromptu-parsed.json"


def clean_line(s: str) -> str:
    s = s.strip()
    s = re.sub(r"^[\s\W\d]+", "", s)
    s = s.replace("\u2019", "'").replace("\u2018", "'").replace("\u201c", '"').replace("\u201d", '"')
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return ""
    if not s[0].isalnum() and s[0] not in '"\'(':
        s = s[1:].strip()
    return s


def iter_doc_text(document: "docx.document.Document") -> list[str]:
    lines: list[str] = []
    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        lines.append(t)
    return lines


def main() -> None:
    pdfs = sorted(REPO.glob("WSDA*.docx"))
    if not pdfs:
        raise SystemExit(f"No WSDA*.docx under {REPO}")
    path = pdfs[0]
    d = docx.Document(str(path))
    raw_lines = iter_doc_text(d)
    seen: set[str] = set()
    motions: list[dict] = []
    n = 0
    for line in raw_lines:
        text = clean_line(line)
        if len(text) < 12:
            continue
        key = re.sub(r"\s+", " ", text.lower())
        if key in seen:
            continue
        seen.add(key)
        text = text.replace("kids'screen", "kids' screen")
        text = re.sub(r"\bAl-generated\b", "AI-generated", text)
        text = re.sub(r"\bAl can\b", "AI can", text)
        n += 1
        motions.append(
            {
                "id": f"wsda-impromptu-{n:03d}",
                "number": n,
                "text": text,
                "format": "WSDA impromptu (from topic bank docx)",
                "sourceName": "WSDA即兴辩论题库 (Word topic bank)",
                "sourceUrl": None,
                "sourceFile": path.name,
            }
        )

    OUT.write_text(
        json.dumps({"version": 1, "motions": motions}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print("Wrote", len(motions), "motions from", path.name, "->", OUT)


if __name__ == "__main__":
    main()
